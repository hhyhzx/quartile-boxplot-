"""
第二十四章《数据的分析》配套练习 — 每个知识点 2 道题
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import os

plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
plt.rcParams['axes.unicode_minus'] = False

BASE = 'E:/maths_work'
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_problem(num, ptype, text):
    p = doc.add_paragraph()
    r = p.add_run(f'第{num}题')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(180, 40, 40)
    r2 = p.add_run(f'  ({ptype})')
    r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(100, 100, 100)
    p2 = doc.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(4)
    return p2

def add_solution(text):
    p = doc.add_paragraph()
    r = p.add_run('解答：')
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0, 100, 0)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(50, 50, 50)
    p.paragraph_format.space_after = Pt(10)

def add_figure(filename, width=3.5):
    path = os.path.join(BASE, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Inches(width))
        p.paragraph_format.space_after = Pt(4)

def add_separator():
    doc.add_paragraph('─' * 50)

# ================================================================
# Generate figures
# ================================================================

# Figure for Problem 5 (quartile/boxplot)
data_q = np.array([3, 7, 8, 10, 12, 15, 16, 18, 20, 22, 25])
q1v, q2v, q3v = np.percentile(data_q, 25), np.percentile(data_q, 50), np.percentile(data_q, 75)
iqr_v = q3v - q1v

fig1, ax = plt.subplots(figsize=(8, 2.8))
ax.set_xlim(0, 28); ax.set_ylim(-0.6, 0.8)
ax.set_title('箱线图 (第5题)', fontsize=13, fontweight='bold', color='#1D3557')
for spine in ax.spines.values(): spine.set_visible(False)
ax.grid(axis='x', alpha=0.2); ax.set_yticks([])

# Box
box = mpatches.FancyBboxPatch((q1v, -0.22), iqr_v, 0.44, boxstyle="round,pad=0.02",
                               facecolor='#A8DADC', edgecolor='#457B9D', lw=2, zorder=3)
ax.add_patch(box)
# Median line
ax.plot([q2v, q2v], [-0.22, 0.22], color='#1D3557', lw=3, zorder=5)
# Whiskers
ax.plot([data_q[0], q1v], [0, 0], color='#457B9D', lw=1.8, zorder=3)
ax.plot([q3v, data_q[-1]], [0, 0], color='#457B9D', lw=1.8, zorder=3)
# Caps
for xv in [data_q[0], data_q[-1]]:
    ax.plot([xv, xv], [-0.1, 0.1], color='#457B9D', lw=1.8, zorder=3)
# Labels
for xv, lbl in [(data_q[0], f'min={data_q[0]}'), (q1v, f'Q1={q1v:.0f}'),
                 (q2v, f'Q2={q2v:.0f}'), (q3v, f'Q3={q3v:.0f}'), (data_q[-1], f'max={data_q[-1]}')]:
    ax.text(xv, -0.48, lbl, ha='center', fontsize=9, fontweight='bold',
            color='#E63946' if 'Q' in lbl else '#333')

ax.text(14, 0.45, f'IQR = Q3-Q1 = {iqr_v:.0f}', ha='center', fontsize=11, color='#E63946', fontweight='bold')
plt.tight_layout()
fig1.savefig(f'{BASE}/exercise_boxplot.png', dpi=150, bbox_inches='tight', facecolor='white')
plt.close(fig1)

# Figure for Problem 8 (histogram)
groups = ['40-50', '50-60', '60-70', '70-80', '80-90']
freqs = [3, 5, 12, 8, 2]
midpoints = np.array([45, 55, 65, 75, 85])

fig2, ax = plt.subplots(figsize=(7, 3.5))
bars = ax.bar(midpoints, freqs, width=9, color='#A8DADC', edgecolor='#457B9D', lw=1.5, alpha=0.85, zorder=3)
ax.plot(midpoints, freqs, 'o-', color='#E63946', lw=2, ms=8, zorder=5, label='频数折线')
for mid, f in zip(midpoints, freqs):
    ax.text(mid, f + 0.2, str(f), ha='center', fontsize=12, fontweight='bold')
ax.set_xticks([40, 50, 60, 70, 80, 90])
ax.set_xlabel('分数段', fontsize=12)
ax.set_ylabel('频数(人数)', fontsize=12)
ax.set_title('成绩分布直方图 (第8题)', fontsize=13, fontweight='bold', color='#1D3557')
ax.legend(fontsize=10)
ax.grid(axis='y', alpha=0.2)
for spine in ax.spines.values(): spine.set_visible(False)
plt.tight_layout()
fig2.savefig(f'{BASE}/exercise_histogram.png', dpi=150, bbox_inches='tight', facecolor='white')
plt.close(fig2)

print('Figures generated.')

# ================================================================
# TITLE
# ================================================================
title = doc.add_heading('第二十四章《数据的分析》· 配套练习', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('4个知识点 × 2题 = 8道练习题（含详细解答）')
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0, 70, 130)
doc.add_paragraph()

# ================================================================
# PART 1: 集中趋势
# ================================================================
add_heading('一、数据的集中趋势（第1–2题）', level=1)

add_problem(1, '平均数/中位数/众数',
    '某班 10 名学生的数学测验成绩如下（单位：分）：\n'
    '    78,  82,  85,  90,  88,  76,  95,  88,  84,  90\n\n'
    '（1）求这组数据的算术平均数；\n'
    '（2）求中位数；\n'
    '（3）求众数。')
add_solution(
    '（1）平均数 = (78+82+85+90+88+76+95+88+84+90) ÷ 10 = 856 ÷ 10 = 85.6（分）\n'
    '（2）排序：76, 78, 82, 84, 85, 88, 88, 90, 90, 95\n'
    '    n=10 为偶数，中位数取第5、6个的平均数 → (85+88)÷2 = 86.5（分）\n'
    '（3）88 和 90 各出现 2 次，均为众数 → 众数为 88 和 90\n\n'
    '思考：平均数85.6 < 中位数86.5，说明偏低的数据（76,78）拉低了平均数。')

add_problem(2, '加权平均数',
    '某同学本学期数学各项成绩如下表：\n\n'
    '  平时作业：88分（权重 15%）\n'
    '  课堂表现：92分（权重 15%）\n'
    '  期中考试：78分（权重 30%）\n'
    '  期末考试：85分（权重 40%）\n\n'
    '求该同学本学期的加权平均分（期末总评）。')
add_solution(
    '加权平均数 = (88×0.15 + 92×0.15 + 78×0.30 + 85×0.40) / (0.15+0.15+0.30+0.40)\n'
    '    = (13.2 + 13.8 + 23.4 + 34.0) / 1.0\n'
    '    = 84.4（分）\n\n'
    '与简单平均数 (88+92+78+85)÷4 = 85.75 不同，\n'
    '期中、期末权重更大，而期中成绩较低(78)，所以加权分数(84.4)略低。\n'
    '这体现了"重要程度不同，贡献也不同"的思想。')

add_separator()

# ================================================================
# PART 2: 离散程度
# ================================================================
add_heading('二、数据的离散程度（第3–4题）', level=1)

add_problem(3, '方差对比',
    '甲、乙两名射击运动员各进行 5 次射击，成绩如下（单位：环）：\n\n'
    '  甲：8, 9, 10, 9, 9\n'
    '  乙：10, 7, 10, 8, 10\n\n'
    '（1）分别求甲、乙的平均数；\n'
    '（2）分别求甲、乙的方差；\n'
    '（3）哪名运动员的成绩更稳定？为什么？')
add_solution(
    '（1）甲平均数 = (8+9+10+9+9)÷5 = 45÷5 = 9（环）\n'
    '    乙平均数 = (10+7+10+8+10)÷5 = 45÷5 = 9（环）\n\n'
    '（2）甲方差：离差分别为 -1, 0, 1, 0, 0\n'
    '    Q_甲 = (-1)²+0²+1²+0²+0² = 2\n'
    '    s²_甲 = 2÷5 = 0.40\n'
    '    乙方差：离差分别为 1, -2, 1, -1, 1\n'
    '    Q_乙 = 1²+(-2)²+1²+(-1)²+1² = 8\n'
    '    s²_乙 = 8÷5 = 1.60\n\n'
    '（3）甲方差(0.40) < 乙方差(1.60)，甲的成绩更稳定。\n'
    '    虽然两人平均成绩相同(9环)，但甲各次成绩集中在9附近，\n'
    '    乙的成绩波动较大（7~10），方差很好地反映了这一差异。')

add_problem(4, '离差平方和/标准差',
    '已知一组数据：70, 75, 80, 85, 90。\n\n'
    '（1）求这组数据的平均数；\n'
    '（2）求离差平方和 Q；\n'
    '（3）求方差 s² 和标准差 s（保留两位小数）。')
add_solution(
    '（1）平均数 = (70+75+80+85+90)÷5 = 400÷5 = 80\n\n'
    '（2）离差：70-80=-10, 75-80=-5, 80-80=0, 85-80=5, 90-80=10\n'
    '    Q = (-10)² + (-5)² + 0² + 5² + 10² = 100+25+0+25+100 = 250\n\n'
    '（3）方差 s² = Q ÷ n = 250 ÷ 5 = 50.00\n'
    '    标准差 s = √50 ≈ 7.07\n\n'
    '含义：这组数据的平均数 80 分，平均偏离程度约 7.07 分。')

add_separator()

# ================================================================
# PART 3: 四分位数与箱线图
# ================================================================
add_heading('三、四分位数与箱线图（第5–6题）★新课标', level=1)

add_problem(5, '四分位数计算+箱线图',
    '将下列 11 个数据从小到大排列：\n\n'
    '    3,  7,  8,  10,  12,  15,  16,  18,  20,  22,  25\n\n'
    '（1）求第一四分位数 Q1、中位数 Q2、第三四分位数 Q3；\n'
    '（2）求四分位距 IQR；\n'
    '（3）判断数据 3 是否为异常值（小于 Q1−1.5×IQR 视为异常）。')
add_figure('exercise_boxplot.png', 3.8)
add_solution(
    '（1）n = 11（奇数）。\n'
    '    Q2 位置 = (11+1)÷2 = 第6个 → Q2 = 15\n'
    '    Q1：前5个数(3,7,8,10,12)的中位数 → Q1 = 8\n'
    '    Q3：后5个数(16,18,20,22,25)的中位数 → Q3 = 20\n\n'
    '（2）IQR = Q3 − Q1 = 20 − 8 = 12\n\n'
    '（3）下界 = Q1 − 1.5×IQR = 8 − 1.5×12 = 8 − 18 = −10\n'
    '    min = 3 > −10，故 3 不是异常值。\n'
    '    若数据中有值小于 −10 或大于 Q3+1.5×IQR=38，才是异常值。')

add_problem(6, '多组箱线图对比',
    '甲、乙两班某次测验成绩的五数概括如下：\n\n'
    '  甲班：min=48, Q1=62, Q2=74, Q3=85, max=96\n'
    '  乙班：min=55, Q1=68, Q2=73, Q3=78, max=92\n\n'
    '（1）哪个班的中位数更高？\n'
    '（2）哪个班的成绩分布更"集中"？从哪些数据可以看出来？\n'
    '（3）哪个班可能存在更多低分学生？')
add_solution(
    '（1）甲班中位数74 > 乙班中位数73，甲班略高（差距很小）。\n\n'
    '（2）乙班更集中。依据：\n'
    '    乙班 IQR = 78−68 = 10，甲班 IQR = 85−62 = 23；\n'
    '    乙班 max−min = 92−55 = 37，甲班 = 96−48 = 48。\n'
    '    乙班的箱体更窄，须线更短，说明中间50%和整体的离散程度都更小。\n\n'
    '（3）甲班可能存在更多低分学生。\n'
    '    甲班 min=48 低于乙班 min=55，且 Q1=62 低于乙班 Q1=68，\n'
    '    说明甲班后25%的学生成绩更差。')

add_separator()

# ================================================================
# PART 4: 数据分组
# ================================================================
add_heading('四、数据的分组（第7–8题）', level=1)

add_problem(7, '频数分布表',
    '以下是某班 30 名学生的数学测验成绩（单位：分）：\n\n'
    '  62, 68, 71, 73, 75, 76, 78, 79, 80, 81,\n'
    '  82, 83, 84, 85, 85, 86, 87, 88, 88, 89,\n'
    '  90, 91, 91, 92, 93, 94, 95, 96, 97, 98\n\n'
    '（1）以 10 分为组距，将数据分组（40~50, 50~60, ...）；\n'
    '（2）列出频数分布表（含频数和频率）；\n'
    '（3）成绩在 80 分及以上的学生占百分之几？')
add_solution(
    '（1）组距=10，分为6组：40~50, 50~60, 60~70, 70~80, 80~90, 90~100\n\n'
    '（2）频数分布表：\n'
    '    40~50: 0人, 频率=0.00\n'
    '    50~60: 0人, 频率=0.00\n'
    '    60~70: 2人(62,68), 频率=2/30≈0.07\n'
    '    70~80: 6人(71,73,75,76,78,79), 频率=6/30=0.20\n'
    '    80~90: 12人, 频率=12/30=0.40\n'
    '    90~100: 10人, 频率=10/30≈0.33\n'
    '    合计: 30人, 频率=1.00\n\n'
    '（3）80分及以上 = 12+10 = 22人, 占 22/30 ≈ 73.3%\n'
    '    大部分同学成绩集中在80~100分段，整体成绩较好。')

add_problem(8, '直方图+加权估计',
    '某班 40 名学生成绩分组统计如下表：\n\n'
    '  分数段  | 40~50 | 50~60 | 60~70 | 70~80 | 80~90 | 90~100\n'
    '  频数    |   1   |   3   |   8   |  15   |  10   |   3\n\n'
    '（1）画出频数分布直方图的大致形状；\n'
    '（2）用组中值估计全班的平均分；\n'
    '（3）频数最多的组是哪个？说明了什么？')
add_figure('exercise_histogram.png', 3.2)
add_solution(
    '（1）直方图见右图，横轴为分数段，纵轴为频数。\n'
    '    频数折线（红色）先上升后下降，呈"中间高、两边低"的形态。\n\n'
    '（2）组中值×频数：45×1+55×3+65×8+75×15+85×10+95×3\n'
    '    = 45+165+520+1125+850+285 = 2990\n'
    '    估计平均分 = 2990÷40 = 74.75（分）\n\n'
    '（3）频数最多的组是 70~80（15人），说明大部分学生成绩在70~80分之间，\n'
    '    分布呈"中间多、两头少"的特点，近似正态分布。')

# ================================================================
# Summary
# ================================================================
doc.add_page_break()
add_heading('参考答案速查', level=1)

table = doc.add_table(rows=9, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['题号', '知识点', '题型', '答案概要']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs: r.bold = True

answers = [
    ['1', '集中趋势', '计算', 'x̄=85.6, 中位数=86.5, 众数=88和90'],
    ['2', '加权平均数', '计算', '加权平均=84.4分'],
    ['3', '方差对比', '计算+分析', '甲方差0.40 < 乙方差1.60, 甲更稳定'],
    ['4', '离差/标准差', '计算', 'Q=250, s²=50, s≈7.07'],
    ['5', '四分位数', '计算+判断', 'Q1=8, Q2=15, Q3=20, IQR=12, 无异常值'],
    ['6', '箱线图对比', '读图分析', '甲班中位数略高, 乙班更集中, 甲班低分更多'],
    ['7', '频数分布', '统计', '80分以上占73.3%'],
    ['8', '直方图', '作图+估算', '估计平均分=74.75, 70~80分组人最多'],
]
for i, row in enumerate(answers):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# Save
output_path = f'{BASE}/统计数据_8道练习题.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
